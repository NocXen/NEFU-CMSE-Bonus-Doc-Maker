import os
import re
import shutil
from pathlib import Path
from docx import Document
from docxcompose.composer import Composer
from docx.enum.text import WD_BREAK

"""
拼接器 v1.1，
DeepSeek R1 制作。
NocXen 修正。
翻新：2次。

Combiner v1.1 By DeepSeek R1.
Fixed By NocXen.
Rebuilt 2 times.

用途：将指定路径的所有docx文档合并为一个文件。
      采用第一个文件作为模板。
      文档文件之间自动加分页符。
      依赖：docxcompose python-docx
"""

# --------------------------------------------------
# 配置和路径设置
# --------------------------------------------------

INPUT_DIR = Path("./Output_输出")
OUTPUT_DIR = Path("./Output_输出/Combined_合并")

# 创建输出目录
OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    
# --------------------------------------------------
# 文档排序
# --------------------------------------------------
def extract_numbers(filename):
    """
    从文件名中提取数字部分用于排序
    例如：院审批单_击掌活动_2_1.docx -> (2, 1)
    """
    # 使用正则表达式匹配文件名中的数字
    numbers = re.findall(r'(\d+)', str(filename))
    if numbers:
        # 将字符串数字转换为整数
        return tuple(int(num) for num in numbers)
    else:
        return (0,)  # 如果没有数字，返回(0,)

def sort_files(files):
    """
    按照文件名中的数字顺序对文件进行排序
    """
    return sorted(files, key=extract_numbers)

def get_files_by_prefix(directory, prefix):
    """
    获取指定目录下以特定前缀开头的所有文件
    """
    path = Path(directory)
    return [f for f in path.iterdir() if f.is_file() and f.name.startswith(prefix)]

# --------------------------------------------------
# 合并文档
# --------------------------------------------------

def merge_word_documents_t(files, output_path):
    """使用docxcompose库合并文档"""
    """暂时性废弃"""
    
    master = Document(files[0])
    composer = Composer(master)
    
    for file_path in files[1:]:
        doc = Document(file_path)
        composer.append(doc)
    
    composer.save(output_path)

def merge_word_documents(files, output_path):
    """使用docxcompose合并文档，并在每个文档前添加分页符"""
    if not files:
        print("没有找到要合并的文件")
        return
    
    # 创建主文档
    master = Document(files[0])
    print(f"已合并: {files[0].name}")
    composer = Composer(master)
    
    for file_path in files[1:]:
        try:
            # 打开源文档
            doc = Document(file_path)
            
            # 在追加新文档前，先给当前文档添加分页符
            composer.doc.add_page_break()
            
            # 追加文档
            composer.append(doc)
            print(f"已合并: {file_path.name}")
            
        except Exception as e:
            print(f"处理文件 {file_path} 时出错: {e}")
            continue
    
    composer.save(output_path)
    print(f"合并完成，保存为: {output_path}\n")

# --------------------------------------------------
# 主程序
# --------------------------------------------------

def main():
    # 检查输入目录是否存在
    if not INPUT_DIR.exists():
        print(f"错误: 输入目录 {INPUT_DIR} 不存在")
        return
    
    # 获取院审批单和院回执单文件
    approval_files = get_files_by_prefix(INPUT_DIR, "院审批单")
    receipt_files = get_files_by_prefix(INPUT_DIR, "院回执单")
    
    print(f"找到院审批单文件: {len(approval_files)} 个")
    print(f"找到院回执单文件: {len(receipt_files)} 个")
    
    # 按数字顺序排序文件
    sorted_approval_files = sort_files(approval_files)
    sorted_receipt_files = sort_files(receipt_files)
    
    # 打印排序后的文件列表（用于验证）
    print("\n院审批单文件排序结果:")
    for f in sorted_approval_files:
        print(f"  {f.name}")
    
    print("\n院回执单文件排序结果:")
    for f in sorted_receipt_files:
        print(f"  {f.name}")
    print('')
    
    # 合并院审批单
    if sorted_approval_files:
        approval_output = OUTPUT_DIR / "院审批单_Combined.docx"
        merge_word_documents(sorted_approval_files, approval_output)
    else:
        print("没有找到院审批单文件，跳过合并")
    
    # 合并院回执单
    if sorted_receipt_files:
        receipt_output = OUTPUT_DIR / "院回执单_Combined.docx"
        merge_word_documents(sorted_receipt_files, receipt_output)
    else:
        print("没有找到院回执单文件，跳过合并")
    
    print(f"所有操作完成！合并后的文件保存在: {OUTPUT_DIR}")

if __name__ == "__main__":
    print('='*25)
    print(f"正在运行 {os.path.splitext(os.path.basename(__file__))[0]} ...")
    print('='*25, "\n")
    main()
    input("\n按回车键退出...\n")
else:
    print('='*25)
    print(f"正在加载 {__name__} ...")
    print('='*25, "\n")