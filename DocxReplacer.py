from docx import Document
import re
import os
from pathlib import Path

"""
文档替换脚本 v2.2，
NocXen 制作。
由 DeepSeek R1 翻新 v2.0。
重大翻新：11次。

DocxReplacer v2.2 By NocXen.
v2.0 Reconstructed By DeepSeek R1.
Rebuilt 11 times.

用途：读取指定数据文件 ( DR_data.txt ) ，
      根据数据文件，寻找模板文档，
      并根据数据文件里的替换字段，
      在文档中寻找相应占位字段，
      进行 直接替换 / 特殊处理后替换，
      替换后，根据数据文件，保存替换后文件。
      依赖：python-docx
"""

# --------------------------------------------------
# 配置和路径设置
# --------------------------------------------------

INPUT_DIR = Path("./Template_模板")
OUTPUT_DIR = Path("./Output_输出")
DATA_FILE = Path("DR_data.txt")

# 创建目录
INPUT_DIR.mkdir(exist_ok=True)
OUTPUT_DIR.mkdir(exist_ok=True)

# --------------------------------------------------
# 数字转换函数
# --------------------------------------------------
def number_to_chinese(n, formal=True):
    """数字转换函数"""
    formal_list = ['零', '壹', '贰', '叁', '肆', '伍', '陆', '柒', '捌', '玖']
    informal_list = ['零', '一', '二', '三', '四', '五', '六', '七', '八', '九']
    
    num_list = formal_list if formal else informal_list
    
    if n == 0:
        return num_list[0]
    elif n < 10:
        return num_list[n]
    elif n == 10:
        return '拾' if formal else '十'
    
    shi = n // 10
    ge = n % 10
    
    if formal:
        shi_str = '拾' if shi == 1 else num_list[shi] + '拾'
    else:
        shi_str = '十' if shi == 1 else num_list[shi] + '十'
    
    ge_str = num_list[ge] if ge > 0 else ''
    
    return shi_str + ge_str

# --------------------------------------------------
# 日期处理函数
# --------------------------------------------------
def format_date(date_str, DotSplit=True):
    """格式化日期字符串"""
    try:
        # 使用正则表达式提取所有连续数字
        numbers = re.findall(r'\d+', date_str)
        
        if len(numbers) < 3:
            raise ValueError("需要包含年、月、日信息")
        
        # 处理年份（支持2位和4位年份）
        year_str = numbers[0]
        if len(year_str) == 2:
            year = int("20" + year_str)  # 假设20xx年
        elif len(year_str) == 4:
            year = int(year_str)
        else:
            raise ValueError("年份格式不正确")
        
        month = int(numbers[1])
        day = int(numbers[2])
        
        # 验证日期合理性
        if not (1 <= month <= 12):
            raise ValueError("月份应在1-12之间")
        if not (1 <= day <= 31):
            raise ValueError("日期应在1-31之间")
        
        return f"{year:04d}.{month}.{day}" if DotSplit else f"{year:04d}年{month}月{day}日"
        
    except Exception as e:
        raise ValueError(f"日期格式错误: {e}")

# --------------------------------------------------
# .split字段处理器
# --------------------------------------------------
class SplitFieldProcessor:
    def __init__(self):
        self.split_info = {}  # 记录每个基字段实际切分的数量
    
    def process_split_field(self, field_name, value, row_data):
        """处理.split字段并记录切分信息"""
        if ".split(" in field_name and field_name.endswith(")"):
            delimiter = re.search(r"\.split\((.+)\)", field_name).group(1)
            base_name = field_name.split(".split(")[0]
            values = [v.strip() for v in value.split(delimiter) if v.strip()]
            
            # 记录切分信息
            self.split_info[base_name] = len(values)
            
            # 为每个分割值创建新字段
            for i, val in enumerate(values, 1):
                new_field = f"{base_name}.split.{i:02d}"
                row_data[new_field] = val
                # [特殊处理: 名字过长预警]
                print(f"! [WARN] 长名字: {val}") if len(val) > 6 and new_field.startswith("RN.split.") else None
            
            return True
        return False
    
    def clean_extra_split_fields(self, doc):
        """清理文档中多余的.split占位符"""
        # 构建清理模式
        patterns = []
        for base_name, actual_count in self.split_info.items():
            # 匹配所有超过实际数量的.split字段
            for i in range(actual_count + 1, 100):  # 假设最多99个
                patterns.append(f"{base_name}.split.{i:02d}")
        
        if not patterns:
            return
        
        # 合并所有模式
        pattern_str = "|".join(re.escape(p) for p in patterns)
        pattern = re.compile(pattern_str)
        
        # 清理段落中的多余占位符
        for para in doc.paragraphs:
            for run in para.runs:
                if pattern.search(run.text):
                    # 替换所有匹配的占位符为空字符串
                    run.text = pattern.sub("", run.text)
        
        # 清理表格中的多余占位符
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        for run in para.runs:
                            if pattern.search(run.text):
                                run.text = pattern.sub("", run.text)

# --------------------------------------------------
# 多行数据解析器
# --------------------------------------------------
class MultiLineParser:
    def __init__(self):
        self.data_rows = []
        self.field_definitions = {}
        
    def parse_input_file(self, filename):
        """解析多行数据文件"""
        with open(filename, 'r', encoding='utf-8') as f:
            lines = f.readlines()
        
        # 提取字段定义
        current_field = None
        field_values = []
        
        for line in lines:
            stripped = line.strip()
            
            # 跳过空行和注释
            if not stripped or stripped.startswith(("#", "--")):
                continue
                
            # 字段定义行
            if stripped.startswith("-  $ "):
                # 保存上一个字段的值
                if current_field and field_values:
                    self.field_definitions[current_field] = field_values
                    field_values = []
                
                # 解析新字段
                field_line = stripped[5:]
                field_name = field_line.split(':', 1)[0].strip()
                field_name = re.sub(r'\[.*?\]', '', field_name).strip()
                current_field = field_name
                
            # 值行
            elif current_field and stripped:
                field_values.append(stripped)
        
        # 保存最后一个字段
        if current_field and field_values:
            self.field_definitions[current_field] = field_values
        
        # 确定数据行数（以Output.files为准）
        output_files = self.field_definitions.get("Output.files", [])
        num_rows = len(output_files)
        
        if num_rows == 0:
            raise ValueError("未找到Output.files字段")
        
        # 为每行创建数据字典
        for i in range(num_rows):
            row_data = {}
            for field, values in self.field_definitions.items():
                # 如果该字段有足够的行，使用对应行的值，否则使用最后一行的值
                if i < len(values):
                    row_data[field] = values[i]
                else:
                    row_data[field] = values[-1]
            
            self.data_rows.append(row_data)
        
        return self.data_rows

# --------------------------------------------------
# 字段处理函数
# --------------------------------------------------
def process_field_value(field_name, value, row_data, split_processor):
    """处理单个字段的值"""
    # 优先处理.split字段
    if split_processor.process_split_field(field_name, value, row_data):
        return None
    
    # 处理日期字段
    if field_name.endswith(".dateD"):
        return format_date(value, DotSplit=True)
    elif field_name.endswith(".date"):
        return format_date(value, DotSplit=False)
    
    # 处理数字字段
    if field_name.endswith(".CL.int"):
        if value.isdigit():
            return number_to_chinese(int(value), formal=True)
    elif field_name.endswith(".CLS.int"):
        if value.isdigit():
            return number_to_chinese(int(value), formal=False)
    
    return value

def build_replacements(row_data, split_processor):
    """为单行数据构建替换字典"""
    replacements = {}
    
    # 先处理基本字段
    for field, value in row_data.items():
        processed = process_field_value(field, value, replacements, split_processor)
        if processed is not None:
            replacements[field] = processed
    
    # 增加占位空格，以增加美观性  [特殊处理]
    # 处理：'Time.Act.Pub.date'
    replacements["Time.Act.Pub.date"] = (lambda month, day: f"{' ' if month < 10 else ''}{' ' if day < 10 else ''}")(*[int(n) for n in re.findall(r'\d+', replacements["Time.Act.Pub.date"])[1:3]]) + replacements["Time.Act.Pub.date"]
    # 处理: 'Reason.DM.Maker'
    replacements["Reason.DM.Maker"] += ' ' * (38 - 2*len(replacements["Reason.DM.Maker"]))
    
    # 处理复合字段，提取时间  [特殊处理]
    activity_date = replacements.get("Time.Act.Pub.dateD", "0000.00.00")
    month_num = int(activity_date.split('.')[1])
    replacements["Time.Act.CLS.int"] = number_to_chinese(month_num, formal=False)
    
    # 优秀部员/部长、加分类型、活动描述处理  [特殊处理]
    em_act_input = row_data.get("Reason.Activity", "").strip()
    em_act_map = {
        "优秀部长": "优秀部长",
        "优秀部员": "优秀部员"
    }
    
    for k, v in em_act_map.items():
        if k in em_act_input:
            replacements["Reason.Activity"] = f"在{replacements['Time.Act.CLS.int']}月份工作中表现积极"
            em_part = f"，被评为{replacements['Time.Act.CLS.int']}月份{v}"
            break
        else:
            replacements["Reason.Activity"] = replacements["Reason.Activity"] if replacements["Reason.Activity"].endswith("，表现突出") and replacements["Reason.Activity"].startswith("参与") else f"参与{replacements.get("Reason.Activity", "")}，表现突出"
            em_part = ""
        
    score_value = int(row_data.get("Score.CL.int", 0))
    score_type = row_data.get("Score.type", "").lower()
    type_part = ""
    
    if score_type == '文体':
        if score_value >= 6:
            type_part = "，给予每人一次院级通报表扬"
        elif score_value >= 4:
            type_part = "，给予每人一次院级点名表扬"
    elif score_type == '德育':
        if score_value >= 3:
            type_part = "，给予每人一次院级通报表扬"
        elif score_value >= 2:
            type_part = "，给予每人一次院级点名表扬"
    
    replacements["Reason.EM&type"] = em_part + type_part
    
    return replacements

# --------------------------------------------------
# 文档处理函数
# --------------------------------------------------
def replace_in_document(template_path, output_path, replacements, split_processor):
    """替换文档中的字段并清理多余占位符"""
    try:
        # 构建完整路径
        full_template_path = INPUT_DIR / template_path
        full_output_path = OUTPUT_DIR / output_path
        
        if not full_template_path.exists():
            print(f"错误: 模板文件不存在 {full_template_path}")
            return False
        
        doc = Document(full_template_path)
        
        # 替换段落中的文本
        for p in doc.paragraphs:
            for run in p.runs:
                for field, value in replacements.items():
                    if field in run.text:
                        run.text = run.text.replace(field, str(value))
        
        # 替换表格中的文本
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        for run in p.runs:
                            for field, value in replacements.items():
                                if field in run.text:
                                    run.text = run.text.replace(field, str(value))
        
        # 清理多余的.split占位符
        split_processor.clean_extra_split_fields(doc)
        
        # 确保输出目录存在
        full_output_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(full_output_path)
        print(f"✓ 已生成: {output_path}")
        return True
        
    except Exception as e:
        print(f"✗ 处理文档出错: {e}")
        return False

# --------------------------------------------------
# 主程序
# --------------------------------------------------
def main():
    if not DATA_FILE.exists():
        print(f"错误: 数据文件不存在 {DATA_FILE}")
        return
    
    try:
        # 解析多行数据
        parser = MultiLineParser()
        data_rows = parser.parse_input_file(DATA_FILE)
        
        print(f"找到 {len(data_rows)} 行数据需要处理")
        
        # 处理每一行数据
        for i, row_data in enumerate(data_rows, 1):
            print(f"\n处理第 {i} 行数据:")
            
            # 创建.split字段处理器
            split_processor = SplitFieldProcessor()
            
            # 构建替换字典
            replacements = build_replacements(row_data, split_processor)
            
            # 获取模板和输出文件
            template_files = row_data.get("Template.files", "").split()
            output_files = row_data.get("Output.files", "").split()
            
            if not template_files:
                print("错误: 未指定模板文件")
                continue
            
            # 生成输出文件序列
            if len(output_files) == len(template_files):
                output_sequence = output_files
            elif len(output_files) == 1:
                base_name, ext = os.path.splitext(output_files[0])
                output_sequence = [output_files[0]] + [f"{base_name}_{j}{ext}" for j in range(1, len(template_files))]
            else:
                output_sequence = [f"output_{i}_{j}.docx" for j in range(len(template_files))]
            
            # 处理每个模板
            for template, output in zip(template_files, output_sequence):
                success = replace_in_document(template, output, replacements, split_processor)
                if not success:
                    print(f"处理失败: {template} -> {output}")
        
        print("\n处理完成！\n")
        
    except Exception as e:
        print(f"程序执行出错: {e}")

if __name__ == "__main__":
    print('='*25)
    print(f"正在运行 {os.path.splitext(os.path.basename(__file__))[0]} ...")
    print('='*25, "\n")
    main()
    input("\n按回车键退出...\n")
else:
    print('='*25)
    print(f"正在加载 {__name__} ...")
    print('='*25, "\n")